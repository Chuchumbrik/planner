"""Convert project UX markdown report to Word .docx (no pandoc required)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import Table


def set_cell_shading(cell, fill: str) -> None:
    """fill: hex without #, e.g. 'F2F2F2'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_formatted_runs(paragraph, text: str) -> None:
    """Support **bold** and *italic* (single asterisk, non-greedy)."""
    if not text:
        return
    # Split preserving **bold** and `code`
    pos = 0
    while pos < len(text):
        if text.startswith("**", pos):
            end = text.find("**", pos + 2)
            if end != -1:
                run = paragraph.add_run(text[pos + 2 : end])
                run.bold = True
                pos = end + 2
                continue
        # italic *...* but not **
        m = re.match(r"\*([^*]+)\*", text[pos:])
        if m and not text[pos : pos + 2] == "**":
            run = paragraph.add_run(m.group(1))
            run.italic = True
            pos += m.end()
            continue
        # find next special
        next_b = text.find("**", pos)
        next_i = text.find("*", pos)
        if next_i == pos and text.startswith("**", pos):
            next_special = next_b
        else:
            candidates = [x for x in (next_b, next_i) if x != -1]
            next_special = min(candidates) if candidates else -1
        if next_special == -1:
            paragraph.add_run(text[pos:])
            break
        if next_special > pos:
            paragraph.add_run(text[pos:next_special])
        pos = next_special


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def is_table_sep(line: str) -> bool:
    s = line.strip().strip("|")
    if not s:
        return False
    parts = [p.strip() for p in s.split("|")]
    return all(re.match(r"^:?-+:?$", p) for p in parts if p)


def parse_table_row(line: str) -> list[str]:
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [c.strip() for c in inner.split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table: Table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row_cells in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            text = row_cells[j] if j < len(row_cells) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, text)
            if i == 0:
                set_cell_shading(cell, "E8E8E8")
                for run in p.runs:
                    run.bold = True


def convert_md_to_docx(md_path: Path, out_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if is_table_row(line):
            table_rows: list[list[str]] = []
            while i < len(lines) and (
                is_table_row(lines[i]) or (lines[i].strip() and is_table_sep(lines[i]))
            ):
                if is_table_sep(lines[i]):
                    i += 1
                    continue
                if is_table_row(lines[i]):
                    table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s*", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, text)
            i += 1
            continue

        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, text)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        # continuation or new paragraph
        buf = [line]
        i += 1
        while i < len(lines):
            nl = lines[i]
            if (
                nl.strip() == ""
                or nl.startswith("#")
                or is_table_row(nl)
                or nl.strip().startswith("- ")
                or re.match(r"^\d+\.\s", nl.strip())
                or nl.strip() == "---"
            ):
                break
            buf.append(nl)
            i += 1
        para_text = "\n".join(buf).strip()
        if para_text.startswith("*") and para_text.endswith("*") and para_text.count("*") == 2:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            add_formatted_runs(p, para_text)
            for run in p.runs:
                run.italic = True
        else:
            for k, chunk in enumerate(para_text.split("\n")):
                if k > 0:
                    doc.add_paragraph()
                p = doc.add_paragraph()
                add_formatted_runs(p, chunk.strip())

    doc.save(out_path)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "UX_ANALYSIS_REPORT.md"
    out = root / "docs" / "UX_ANALYSIS_REPORT.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    if not md.is_file():
        print(f"Missing: {md}", file=sys.stderr)
        return 1
    convert_md_to_docx(md, out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
